"""
Multi-format document loader.

Design choice: return one Document per page (PDF) or section (DOCX)
rather than one giant Document per file. This preserves structural
metadata (page numbers, section headings) that becomes searchable
in the retriever and displayable in citations.
"""
import logging
from pathlib import Path
from typing import List, Union
from docmind.models import Document

logger = logging.getLogger(__name__)


def load_pdf(path: Path) -> List[Document]:
    """Extract text page-by-page using pdfplumber.
    pdfplumber handles multi-column layouts and tables better than PyPDF2."""
    import pdfplumber

    docs = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(Document(
                    content=text.strip(),
                    source=str(path.resolve()),
                    metadata={
                        "file_type": "pdf",
                        "page": page_num + 1,
                        "total_pages": total,
                        "filename": path.name,
                    }
                ))

    logger.info(f"PDF '{path.name}': loaded {len(docs)}/{total} non-empty pages")
    return docs


def load_docx(path: Path) -> List[Document]:
    """Load DOCX preserving heading-based section structure.
    Each top-level section becomes a Document, with its heading in metadata."""
    from docx import Document as DocxDoc

    docx = DocxDoc(path)
    sections: List[tuple] = []
    current_heading = "Preamble"
    current_lines: List[str] = []

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current_lines:
                sections.append((current_heading, "\n".join(current_lines)))
                current_lines = []
            current_heading = text
        else:
            current_lines.append(text)

    if current_lines:
        sections.append((current_heading, "\n".join(current_lines)))

    docs = [
        Document(
            content=content,
            source=str(path.resolve()),
            metadata={
                "file_type": "docx",
                "section": heading,
                "filename": path.name,
            }
        )
        for heading, content in sections
        if content.strip()
    ]

    logger.info(f"DOCX '{path.name}': loaded {len(docs)} sections")
    return docs


def load_text(path: Path) -> List[Document]:
    """Load plain text or markdown as a single Document."""
    content = path.read_text(encoding="utf-8", errors="replace").strip()
    file_type = "markdown" if path.suffix.lower() in (".md", ".markdown") else "text"
    return [Document(
        content=content,
        source=str(path.resolve()),
        metadata={"file_type": file_type, "filename": path.name}
    )]


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,
    ".txt": load_text,
    ".md": load_text,
    ".markdown": load_text,
}


def load_document(path: Union[str, Path]) -> List[Document]:
    """Route to the correct loader based on file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Not found: {path}")
    loader = LOADERS.get(path.suffix.lower())
    if not loader:
        raise ValueError(f"Unsupported extension: {path.suffix}")
    return loader(path)


def load_directory(directory: Union[str, Path]) -> List[Document]:
    """Recursively load all supported files from a directory."""
    directory = Path(directory)
    docs = []
    for file_path in sorted(directory.rglob("*")):
        if file_path.suffix.lower() in LOADERS:
            try:
                docs.extend(load_document(file_path))
            except Exception as exc:
                logger.warning(f"Skipping {file_path.name}: {exc}")
    logger.info(f"Directory '{directory}': loaded {len(docs)} total documents")
    return docs